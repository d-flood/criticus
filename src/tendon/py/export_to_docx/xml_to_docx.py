import os
from pathlib import Path
import re
from typing import List

from lxml import etree as et
from docx import Document
from natsort import natsorted
import PySimpleGUI as sg

from tendon.py.reformat_collation.itsee_to_open_cbgm import reformat_xml
import tendon.py.edit_settings as es

TEI_NS = '{http://www.tei-c.org/ns/1.0}'
XML_NS = '{http://www.w3.org/XML/1998/namespace}'
ABBR_TO_FULL = {
    'Matt': 'Matthew',
    'B01': 'Matthew',
    'Mark': 'Mark',
    'B02': 'Mark',
    'Luke': 'Luke',
    'B03': 'Luke',
    'John': 'John',
    'B04': 'John',
    'Acts': 'Acts',
    'B05': 'Acts',
    'Rom': 'Romans',
    'B06': 'Romans',
    'Romans': 'Romans',
    'R': 'Romans',
    '1 Cor': '1 Corinthians',
    '1Cor': '1 Corinthians',
    'ICor': '1 Corinthians',
    'B07': '1 Corinthians',
    'B07': '1 Corinthians',
    '1 Corinthians': '1 Corinthians',
    '2 Cor': '2 Corinthians',
    '2Cor': '2 Corinthians',
    'IICor': '2 Corinthians',
    '2 Corinthians': '2 Corinthians',
    'B08': '2 Corinthians',
    'Gal': 'Galatians',
    'Galatians': 'Galatians',
    'B09': 'Galatians',
    'Eph': 'Ephesians',
    'Ephesians': 'Ephesians',
    'B10': 'Ephesians',
    'Phil': 'Philippians',
    'Philippians': 'Philippians',
    'B11': 'Philippians',
    'Col': 'Colossians',
    'Colossians': 'Colossians',
    'B12': 'Colossians',
    '1 Thess': '1 Thessalonians',
    '1Thess': '1 Thessalonians',
    '1 Thessalonians': '1 Thessalonians',
    'B13': '1 Thessalonians',
    '2 Thess': '2 Thessalonians',
    '2Thess': '2 Thessalonians',
    '2 Thessalonians': '2 Thessalonians',
    'B14': '2 Thessalonians',
    '1 Tim': '1 Timothy',
    '1Tim': '1 Timothy',
    '1 Timothy': '1 Timothy',
    'B15': '1 Timothy',
    '2 Tim': '2 Timothy',
    '2Tim': '2 Timothy',
    '2 Timothy': '2 Timothy',
    'B16': '2 Timothy',
    'Titus': 'Titus',
    'B17': 'Titus',
    'Phlm': 'B18',
    'Philemon': 'Philemon',
    'B18': 'Philemon',
    'Heb': 'B19',
    'Hebrews': 'Hebrews',
    'B19': 'Hebrews',
    'Jas': 'James',
    'James': 'James',
    'B20': 'James',
    '1 Pet': '1 Peter',
    '1Pet': '1 Peter',
    '1 Peter': '1 Peter',
    'B21': '1 Peter',
    '2 Pet': '2 Peter',
    '2Pet': '2 Peter',
    '2 Peter': '2 Peter',
    'B22': '2 Peter',
    '1 John': '1 John',
    'B23': '1 John',
    '1John': '1 John',
    '2 John': '2 John',
    'B24': '2 John',
    '2John': '2 John',
    '3 John': '3John',
    '3John': '3John',
    'B25': '3John',
    'Jude': 'Jude',
    'B26': 'Jude',
    'Rev': 'Revelation',
    'Revelation': 'Revelation',
    'B27': 'Revelation',
}

def get_xml_file(xml: str) -> et._Element:
    temp_cx_file = 'temp_xml_collation_file'
    xml = xml.replace('xml:id="1', 'xml:id="I')
    xml = xml.replace('xml:id="2', 'xml:id="II')
    xml = xml.replace('xml:id="3', 'xml:id="III')
    xml = xml.replace('subreading', 'subr')
    with open(temp_cx_file, 'w', encoding='utf-8') as file:
        file.write(xml)
    if re.search('<teiHeader>', xml) is None:
        try:
            temp_cx_file = reformat_xml(temp_cx_file)
        except:
            return None
    parser = et.XMLParser(remove_blank_text=True, encoding='UTF-8')
    tree = et.parse(temp_cx_file, parser) #type: et._ElementTree
    root = tree.getroot()
    os.remove(temp_cx_file)
    return root

def get_document():
    this_dir = Path(__file__).parent
    template = this_dir.joinpath('template.docx').as_posix()
    print(template)
    return Document(template)

def load_xml_file(xml_file: str):
    with open(xml_file, 'r', encoding='utf-8') as file:
        xml = file.read()
    return get_xml_file(xml)

def construct_full_ref(ab: et. _Element):
    ref = ab.get(f'{XML_NS}id').replace('-APP', '') #type: str
    if ref.startswith('B'): # then it is an INTF/IGNTP style reference... probably
        book = re.search(r'B\d+', ref).group(0)
        book = ABBR_TO_FULL[book]
        chapter = re.search(r'K\d+', ref).group(0)
        verse = re.search(r'V\d+', ref).group(0)
        ref = f'{book} {chapter}:{verse}'
    else:
        book = re.search(r'.[a-zA-Z]+', ref)
        if not book:
            return ref
        book = book.group(0)
        full_book = ABBR_TO_FULL.get(book)
        if not full_book:
            return ref
        reference = ref.replace(book, '').replace('.', ':')
        ref = f'{full_book} {reference}'
    return ref

def print_reference(document: Document, ab: et._Element):
    ref = construct_full_ref(ab)
    reference = document.add_paragraph(ref)
    reference.style = document.styles['reference']

def group_basetext_words(basetext: str, words_per_line: int) -> List[list]:
    words_per_line = words_per_line - 1
    grouped_basetext = []
    current_group = []
    chunk = 0
    for word in basetext.split():
        if chunk == words_per_line:
            current_group.append(word)
            grouped_basetext.append(current_group)
            chunk = 0
            current_group = []
            continue
        current_group.append(word)
        chunk += 1
    if current_group != []:
        grouped_basetext.append(current_group)
    return grouped_basetext

def construct_basetext(ab: et._Element) -> str:
    basetext = []
    for elem in ab:
        if elem.tag == f'{TEI_NS}seg':
            basetext.append(elem.text)
        elif elem.tag == f'{TEI_NS}app' and elem.find(f'{TEI_NS}lem').get('type') != 'om':
            basetext.append(elem.find(f'{TEI_NS}lem').text)
    return ' '.join(basetext)

def print_basetext(document: Document, ab: et._Element, words_per_line: int):
    basetext = construct_basetext(ab)
    basetext = group_basetext_words(basetext, words_per_line)
    table = document.add_table(rows=0, cols=10)
    index = 2
    for line in basetext:
        row_cells = table.add_row().cells
        for cell, word in enumerate(line):
            row_cells[cell].text = f"{word}\n{index}"
            row_cells[cell].paragraphs[0].style = document.styles['table cell']
            index += 2

def print_app(document: Document, app: et._Element):
    app_from = app.get('from')
    app_to = app.get('to')
    if app_from == app_to:
        index = app_from
    else:
        index = f'{app_from}–{app_to}'
    p = document.add_paragraph(index)
    p.style = document.styles['index']

def sort_by_ga(wits: List[str]):
    papyri = []
    majuscules = []
    minuscules = []
    lectionaries = []
    editions = []
    for wit in wits:
        if wit.lower().startswith('p'):
            papyri.append(wit)
        elif wit.startswith('0'):
            majuscules.append(wit)
        elif wit[0].isdigit():
            minuscules.append(wit)
        elif wit.lower().startswith('l'):
            lectionaries.append(wit)
        else:
            editions.append(wit)
    return natsorted(papyri) + natsorted(majuscules) + natsorted(minuscules) + natsorted(lectionaries) + natsorted(editions)

def print_rdg(
    document, rdg: et._Element, 
    text_wits_separator: str, 
    rdg_n_text_separator: str, 
    text_bold: bool
    ):
    if rdg.text:
        greek_text = rdg.text
    else:
        greek_text = rdg.get('type')
    p = document.add_paragraph()
    p.style = document.styles['reading']
    rdg_name = re.sub(r'\d', '', rdg.get('n'))
    p.add_run(rdg_name).italic = True
    p.add_run(rdg_n_text_separator)
    p.add_run(greek_text).bold = text_bold
    wits = rdg.get('wit').split(' ')
    wits = sort_by_ga(wits)
    wits = ' '.join(wits)
    p.add_run(f"{text_wits_separator}{wits}")

def save_docx(document: Document, settings: dict):
    docx_filename = sg.popup_get_file('', no_window=True, save_as=True, initial_folder=settings.get('export_docx_folder'), file_types=(('DOCX Files', '*.docx'),))
    if not docx_filename:
        return
    if not docx_filename.endswith('.docx'):
        docx_filename = f'{docx_filename}.docx'
    docx_dir = Path(docx_filename).parent.as_posix()
    es.edit_settings('export_docx_folder', docx_dir)
    document.save(docx_filename)
    return docx_filename

def export_xml_to_docx(xml_filename: str):
    settings = es.get_settings()
    document = get_document()
    root = load_xml_file(xml_filename)
    for ab in root.findall(f'{TEI_NS}ab'):
        print_reference(document, ab)
        print_basetext(document, ab, settings['words_per_line'])
        for app in ab.findall(f'{TEI_NS}app'):
            print_app(document, app)
            for rdg in app.findall(f'{TEI_NS}rdg'): #type: List[et._Element]
                print_rdg(
                    document, rdg, settings['text_wits_separator'], 
                    settings['rdg_n_text_separator'], settings['text_bold']
                    )
    return save_docx(document, settings)
